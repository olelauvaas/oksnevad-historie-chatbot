import streamlit as st
from openai import OpenAI
from PIL import Image
import requests
from io import BytesIO
import tempfile
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from docx import Document
import random

# 🔐 OpenAI API-klient
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

# 🎨 Streamlit-oppsett
st.set_page_config(page_title="Sofies Tidsmaskin")
logo_path = "logo.PNG"
st.image(logo_path, width=300, use_container_width=False)
st.markdown("## Sofies Tidsmaskin")

st.markdown("""
### 🧱 Velkommen til Sofies tidsmaskin
Du er på vei tilbake i tid.

Sammen med Sofie skal du besøke et sted i verden på et bestemt tidspunkt i historien. Der møter du en ungdom som forteller deg hvordan det er å leve akkurat der og da – med sine egne ord.

🔹 Det du skal gjøre:
Skriv inn:
- Navnet ditt
- Dato (f.eks. 18.08.1894)
- Sted og land (f.eks. Bridgetown, Barbados)
- (Valgfritt) etnisitet og samfunnslag
- Velg kjønn på personen du møter

Les historien nøye.
Du får møte en ungdom som forteller om sitt liv, utfordringer, drømmer og samfunnet rundt seg.

Etterpå skal du tenke over:
- Hva lærte du om perioden?
- Hvilke spor av historie kjenner du igjen?
- Hva overrasket deg?

🧠 Husk:
– Det du leser, er en fiksjon basert på historiske forhold.
– Fortellingen er skrevet for å hjelpe deg å forstå tiden gjennom menneskene som levde i den.
– Ingen magi. Ingen roboter. Bare ekte mennesker, ekte liv – slik det kunne ha vært.
""")

# 📦 Session state for historikk
if "story_data" not in st.session_state:
    st.session_state.story_data = {}

# 📸 Bildegenerering

def generer_bilde(prompt):
    try:
        image_response = client.images.generate(
            model="dall-e-3",
            prompt=prompt,
            size="1024x1024",
            quality="standard",
            n=1
        )
        image_url = image_response.data[0].url
        image = Image.open(BytesIO(requests.get(image_url).content))
        return image
    except Exception as e:
        st.error("Kunne ikke generere bilde: " + str(e))
        return None

# 🧾 Brukerinput
if "historie_generert" not in st.session_state or not st.session_state.historie_generert:
    navn = st.text_input("Hva heter du?", placeholder="f.eks. Sofie")
    date = st.text_input("Dato (DD.MM.ÅÅÅÅ)", placeholder="f.eks. 01.05.1897")
    location = st.text_input("Sted og land", placeholder="f.eks. Bridgetown, Barbados")
    extra_details = st.text_input("(Valgfritt) Etnisitet og samfunnslag", placeholder="f.eks. afro-karibisk, arbeiderklasse")
    gender = st.selectbox("Velg kjønn på personen du møter", options=["Tilfeldig", "Jente", "Gutt"], index=0)

    if st.button("Reis i tid") and navn and date and location:

        story_prompt = f"""
🛠️ Systemprompt til GPT: Sofies tidsmaskin

Du er en historiefortellende GPT kalt Sofies tidsmaskin. Brukeren har skrevet inn sitt navn, en dato, et årstall, et sted og et land. Du skal nå ta med brukeren og Sofie (en fiktiv kvinnelig tidsreisepartner) tilbake i tid til dette stedet og tidspunktet.

🎭 Rollen din:
Når dere ankommer, blir dere møtt av en lokal ungdom, som har fått et realistisk navn og kjønn. Personen er den som forteller historien. Personen skal:

- Henvende seg direkte til både Sofie og {navn} i åpningsreplikken.
- Presentere seg med navn og alder – velg et navn som er realistisk for tid, sted og kjønn.
- Du *må* bruke det kjønnet som brukeren har valgt: {gender.lower()}.
- Velg et *realistisk og tydelig kjønnet navn* som var vanlig for {gender.lower()}-barn i {location} i året {date[-4:]}. Bruk kjente navnedatabaser hvis du er usikker.
- Ikke bruk navn som kan forveksles med motsatt kjønn, som "Bram" for jenter eller "Jansje" for gutter.
- Start historien med: "Hei, Sofie og {navn}. Jeg heter [NAVN], og jeg er [ALDER] år."
- Dersom etnisitet og samfunnslag ikke er angitt av brukeren, skal du selv velge og nevne dette tidlig i historien på en naturlig måte.
- Snakke i jeg-form og fortelle en personlig og levende historie om hvordan det er å leve akkurat her og nå.

📜 Historien skal:
- Være **kort og konsis** (maks 500–600 ord), og egnet for ungdom i alderen 16–18 år.
- Inneholde en **drivende konflikt eller dramatisk hendelse** – noe som overrasker eller utfordrer hovedpersonen.
- Ha en tydelig **"wow-faktor"** – noe som gjør at leseren tenker: *"Hæ?! Skjedde DET?!"*
- Inneholde **realistiske og sanselige detaljer** fra tid og sted: arbeid, skole (bare hvis realistisk), familie, samfunn, kultur, politikk.
- Ha en ungdommelig fortellerstil: direkte, ekte og følelsesnær – **unngå overdreven poesi og lange metaforer**.
- Personen takker til slutt Sofie og brukeren for besøket.

🗄 Viktige regler:
- Sofie snakker ikke – hun er bare med på reisen.
- Ikke forklar, oppsummer eller si "Her kommer en historie om...". Gå rett inn i fortellingen med personens første replikk.
- Ikke bruk moderne ord, uttrykk eller konsepter som ikke fantes i perioden (f.eks. plast, strøm, dorullskip).
- Ikke referer til Øksnevad vgs eller andre moderne institusjoner.
Historien foregår i {location} den {date}.
"""

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": story_prompt},
                {"role": "user", "content": "Fortell historien."}
            ]
        )

        story = response.choices[0].message.content

        st.session_state.story_data = {
            "story": story,
            "navn": navn,
            "date": date,
            "location": location,
            "extra_details": extra_details,
            "gender": gender,
            "image": None
        }

        st.session_state.historie_generert = True
        st.rerun()

# ⬇️ Viser historien dersom den er generert
if st.session_state.get("historie_generert"):
    if st.button("🔄 Start på nytt"):
        st.session_state.historie_generert = False
        st.session_state.story_data = {}
        st.rerun()

    st.markdown("---")
    st.markdown("### ✨ Historien fra fortiden")
    st.write(st.session_state.story_data["story"])

    if st.session_state.story_data.get("image") is None:
        bildeprompt = f"A realistic painting of a {st.session_state.story_data['gender'].lower()} teenager in {st.session_state.story_data['location']} in the year {st.session_state.story_data['date'][-4:]}, historical clothing, natural light, facing forward"
        bilde = generer_bilde(bildeprompt)
        st.session_state.story_data["image"] = bilde
        st.rerun()

    if st.session_state.story_data.get("image"):
        st.image(st.session_state.story_data["image"], caption="Historisk portrett")

    st.markdown("""
### 📘 Refleksjonsspørsmål

🗞 Refleksjon etter tidsreisen med Sofie

📍 Ditt valg:
Navn: ___________________________
Dato du besøkte: ___________________
Sted og land: ______________________
Navnet på personen du møtte: ___________________

🔎 1. Hva lærte du?
Svar:

⚡ 2. Hva overrasket deg mest?
Svar:

💬 3. Hva ville du spurt personen om, hvis du fikk stille ett spørsmål?
Svar:

💡 4. Hva kan vi lære av denne tiden i dag?
Svar:

🎯 5. Tidskapsel-score
☑ 1
☑ 2
☑ 3
☑ 4
☑ 5

🧠 Ekstra (valgfritt):
Sammenlign det livet du møtte med ditt eget.
Skriv en kort melding til personen du møtte.
""")

    if st.button("Last ned som Word-dokument"):
        doc = Document()
        doc.add_heading("Historien fra fortiden", 0)
        doc.add_paragraph(st.session_state.story_data["story"])
        doc.add_page_break()
        doc.add_heading("Refleksjonsspørsmål", level=1)
        reflection_text = """
🗞 Refleksjon etter tidsreisen med Sofie

📍 Ditt valg:
Navn: ___________________________
Dato du besøkte: ___________________
Sted og land: ______________________
Navnet på personen du møtte: ___________________

🔎 1. Hva lærte du?
Svar:

⚡ 2. Hva overrasket deg mest?
Svar:

💬 3. Hva ville du spurt personen om?
Svar:

💡 4. Hva kan vi lære av denne tiden i dag?
Svar:

🎯 5. Tidskapsel-score
☑ 1
☑ 2
☑ 3
☑ 4
☑ 5

🧠 Ekstra (valgfritt):
Sammenlign det livet du møtte med ditt eget.
Skriv en kort melding til personen du møtte.
"""
        doc.add_paragraph(reflection_text)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp.seek(0)
            st.download_button(
                label="📄 Last ned Word-dokument",
                data=tmp.read(),
                file_name="sofies_tidsreise.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    if st.button("🔙 Tilbake til start"):
        st.session_state.historie_generert = False
        st.session_state.story_data = {}
        st.rerun()
